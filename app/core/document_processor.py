"""
Document processing service.
Handles multi-format document parsing and text extraction.
"""

import os
from pathlib import Path
from typing import Dict, List, Tuple
import logging

from app.config import get_settings
from app.utils.logger import setup_logger
from app.utils.errors import DocumentProcessingError

logger = setup_logger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Handles document processing for multiple formats."""

    def __init__(self):
        """Initialize document processor."""
        self.supported_formats = settings.supported_formats_list
        logger.info(f"DocumentProcessor initialized. Supported formats: {self.supported_formats}")

    def process_document(self, file_path: str) -> Dict[str, any]:
        """
        Process a document and extract text.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            dict: Processed document data with metadata
            
        Raises:
            DocumentProcessingError: If processing fails
        """
        try:
            # Get file extension
            file_ext = Path(file_path).suffix.lower().strip(".")
            
            if file_ext not in self.supported_formats:
                raise DocumentProcessingError(
                    f"Unsupported file format: {file_ext}"
                )
            
            logger.info(f"Processing document: {file_path} (format: {file_ext})")
            
            # Route to appropriate processor
            if file_ext == "pdf":
                content = self._process_pdf(file_path)
            elif file_ext == "txt":
                content = self._process_txt(file_path)
            elif file_ext == "csv":
                content = self._process_csv(file_path)
            elif file_ext in ["xlsx", "xls"]:
                content = self._process_excel(file_path)
            elif file_ext == "docx":
                content = self._process_docx(file_path)
            elif file_ext == "json":
                content = self._process_json(file_path)
            elif file_ext == "yaml" or file_ext == "yml":
                content = self._process_yaml(file_path)
            else:
                raise DocumentProcessingError(f"Unsupported format: {file_ext}")
            
            # Clean content
            cleaned_content = self._clean_text(content)
            
            logger.info(f"Document processed successfully. Content length: {len(cleaned_content)}")
            
            return {
                "original_path": file_path,
                "file_format": file_ext,
                "content": cleaned_content,
                "content_length": len(cleaned_content),
                "word_count": len(cleaned_content.split()),
                "status": "processed",
            }
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=e)
            raise DocumentProcessingError(f"Failed to process document: {str(e)}")

    def _process_pdf(self, file_path: str) -> str:
        """
        Process PDF file and extract text.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        try:
            from PyPDF2 import PdfReader
            
            text = []
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                logger.info(f"PDF has {len(reader.pages)} pages")
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {str(e)}")
            
            content = "\n".join(text)
            logger.info(f"Extracted {len(content)} characters from PDF")
            return content
            
        except ImportError:
            raise DocumentProcessingError("PyPDF2 library not installed. Install with: pip install PyPDF2")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process PDF: {str(e)}")

    def _process_txt(self, file_path: str) -> str:
        """
        Process TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            str: File content
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            logger.info(f"Extracted {len(content)} characters from TXT")
            return content
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process TXT: {str(e)}")

    def _process_csv(self, file_path: str) -> str:
        """
        Process CSV file and convert to text format.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            str: Formatted table content
        """
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert to formatted string
            content = f"CSV File with {len(df)} rows and {len(df.columns)} columns\n"
            content += f"Columns: {', '.join(df.columns)}\n\n"
            content += df.to_string()
            
            logger.info(f"Extracted {len(df)} rows from CSV")
            return content
            
        except ImportError:
            raise DocumentProcessingError("pandas library not installed. Install with: pip install pandas")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process CSV: {str(e)}")

    def _process_excel(self, file_path: str) -> str:
        """
        Process Excel file and convert to text format.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            str: Formatted sheet content
        """
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets = excel_file.sheet_names
            
            content = f"Excel file with {len(sheets)} sheets\n"
            content += f"Sheets: {', '.join(sheets)}\n\n"
            
            # Extract content from each sheet
            for sheet_name in sheets:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                content += f"\n--- Sheet: {sheet_name} ({len(df)} rows) ---\n"
                content += f"Columns: {', '.join(df.columns)}\n"
                content += df.to_string()
                content += "\n"
            
            logger.info(f"Extracted {len(sheets)} sheets from Excel")
            return content
            
        except ImportError:
            raise DocumentProcessingError("pandas library not installed. Install with: pip install pandas openpyxl")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process Excel: {str(e)}")

    def _process_docx(self, file_path: str) -> str:
        """
        Process DOCX file and extract text.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            str: Extracted text content
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text.append(" | ".join(row_text))
            
            content = "\n".join(text)
            logger.info(f"Extracted {len(content)} characters from DOCX")
            return content
            
        except ImportError:
            raise DocumentProcessingError("python-docx library not installed. Install with: pip install python-docx")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process DOCX: {str(e)}")

    def _process_json(self, file_path: str) -> str:
        """
        Process JSON file and convert to formatted text.
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            str: Formatted JSON content
        """
        try:
            import json
            
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            
            # Convert to formatted string
            content = json.dumps(data, indent=2)
            logger.info(f"Extracted JSON with {len(str(data))} characters")
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process JSON: {str(e)}")

    def _process_yaml(self, file_path: str) -> str:
        """
        Process YAML file and convert to formatted text.
        
        Args:
            file_path: Path to YAML file
            
        Returns:
            str: Formatted YAML content
        """
        try:
            import yaml
            
            with open(file_path, "r", encoding="utf-8") as f:
                data = yaml.safe_load(f)
            
            # Convert back to YAML format
            content = yaml.dump(data, default_flow_style=False)
            logger.info(f"Extracted YAML with {len(str(data))} characters")
            return content
            
        except ImportError:
            raise DocumentProcessingError("PyYAML library not installed. Install with: pip install PyYAML")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to process YAML: {str(e)}")

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.
        
        Args:
            text: Raw text content
            
        Returns:
            str: Cleaned text
        """
        try:
            import re
            
            # Remove extra whitespace
            text = " ".join(text.split())
            
            # Remove control characters
            text = re.sub(r"[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]", "", text)
            
            # Normalize line breaks
            text = text.replace("\r\n", "\n").replace("\r", "\n")
            
            # Remove excessive blank lines
            text = re.sub(r"\n\n+", "\n\n", text)
            
            logger.debug(f"Text cleaned. Original length: {len(text)}")
            return text
            
        except Exception as e:
            logger.warning(f"Error cleaning text: {str(e)}")
            return text

    def extract_metadata(self, file_path: str) -> Dict[str, any]:
        """
        Extract metadata from document.
        
        Args:
            file_path: Path to document
            
        Returns:
            dict: Document metadata
        """
        try:
            file_stat = os.stat(file_path)
            
            metadata = {
                "file_name": os.path.basename(file_path),
                "file_size_bytes": file_stat.st_size,
                "file_size_mb": file_stat.st_size / (1024 * 1024),
                "created_timestamp": file_stat.st_ctime,
                "modified_timestamp": file_stat.st_mtime,
                "file_extension": Path(file_path).suffix.lower().strip("."),
            }
            
            return metadata
            
        except Exception as e:
            logger.warning(f"Error extracting metadata: {str(e)}")
            return {}
